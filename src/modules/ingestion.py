"""
Enhanced document ingestion module with robust PDF processing
Supports multiple extraction methods and advanced text cleaning
"""

import os
import re
import logging
import unicodedata
from typing import Dict, List, Optional, Tuple
from pathlib import Path
from datetime import datetime
from enum import Enum

# PDF Processing
import PyPDF2
import pdfplumber
from pdf2image import convert_from_path
import pytesseract

# Text processing
from langdetect import detect, LangDetectException

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class PDFExtractionMethod(Enum):
    """Available PDF extraction methods"""
    PYPDF2 = "pypdf2"
    PDFPLUMBER = "pdfplumber"
    OCR = "ocr"
    HYBRID = "hybrid"


class TextCleaner:
    """Advanced text cleaning utilities"""
    
    def __init__(self, aggressive: bool = False):
        self.aggressive = aggressive
    
    def remove_excessive_whitespace(self, text: str) -> str:
        """Remove excessive whitespace"""
        text = re.sub(r' +', ' ', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = '\n'.join(line.strip() for line in text.split('\n'))
        return text.strip()
    
    def fix_encoding_issues(self, text: str) -> str:
        """Fix common encoding problems"""
        text = unicodedata.normalize('NFKC', text)
        
        replacements = {
            'â€™': "'", 'â€œ': '"', 'â€': '"',
            'â€"': '—', 'â€"': '–',
            'Ã©': 'é', 'Ã¨': 'è', 'Ã ': 'à',
            'Ã§': 'ç', 'Ã´': 'ô', 'Ã»': 'û',
            'Ã®': 'î', 'Ã«': 'ë',
            '\x00': '', '\ufeff': '',
        }
        
        for old, new in replacements.items():
            text = text.replace(old, new)
        
        return text
    
    def remove_headers_footers(self, text: str, patterns: List[str] = None) -> str:
        """Remove repetitive headers and footers"""
        if patterns is None:
            patterns = [
                r'Page \d+ of \d+',
                r'Page \d+/\d+',
                r'^\d+$',
                r'^Copyright ©.*$',
                r'^Confidential.*$',
            ]
        
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line_stripped = line.strip()
            if not line_stripped:
                cleaned_lines.append(line)
                continue
            
            is_header_footer = any(
                re.search(pattern, line_stripped, re.IGNORECASE) 
                for pattern in patterns
            )
            
            if not is_header_footer:
                cleaned_lines.append(line)
        
        return '\n'.join(cleaned_lines)
    
    def fix_hyphenation(self, text: str) -> str:
        """Fix hyphenated words at line breaks"""
        text = re.sub(r'(\w+)-\s*\n\s*(\w+)', r'\1\2', text)
        return text
    
    def remove_urls_emails(self, text: str) -> str:
        """Remove or normalize URLs and emails"""
        text = re.sub(
            r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+', 
            '[URL]', 
            text
        )
        text = re.sub(
            r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', 
            '[EMAIL]', 
            text
        )
        return text
    
    def normalize_numbers(self, text: str) -> str:
        """Normalize number formats"""
        text = re.sub(r'(\d+),(\d{3})', r'\1\2', text)
        return text
    
    def remove_page_breaks(self, text: str) -> str:
        """Remove page break characters"""
        text = re.sub(r'\f', '\n', text)
        return text
    
    def clean(self, text: str) -> str:
        """Apply complete cleaning pipeline"""
        logger.debug(f"Cleaning text (length: {len(text)})")
        
        text = self.fix_encoding_issues(text)
        text = self.remove_page_breaks(text)
        text = self.fix_hyphenation(text)
        text = self.remove_excessive_whitespace(text)
        text = self.remove_headers_footers(text)
        
        if self.aggressive:
            text = self.remove_urls_emails(text)
        
        text = self.normalize_numbers(text)
        
        logger.debug(f"Text cleaned (final length: {len(text)})")
        return text


class PDFExtractor:
    """Robust PDF text extraction with multiple methods"""
    
    def __init__(self, cleaner: Optional[TextCleaner] = None):
        self.cleaner = cleaner or TextCleaner()
    
    def extract_with_pypdf2(self, path: Path) -> Tuple[str, Dict]:
        """Extract using PyPDF2 (fast but basic)"""
        logger.info(f"[PyPDF2] Extracting: {path.name}")
        text = ""
        metadata = {}
        
        try:
            with open(path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                
                if reader.metadata:
                    metadata = {
                        'title': reader.metadata.get('/Title', ''),
                        'author': reader.metadata.get('/Author', ''),
                        'creation_date': reader.metadata.get('/CreationDate', ''),
                        'num_pages': len(reader.pages)
                    }
                else:
                    metadata = {'num_pages': len(reader.pages)}
                
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
                
                logger.info(f"[PyPDF2] Extracted {len(text)} chars from {metadata.get('num_pages', 0)} pages")
                    
        except Exception as e:
            logger.error(f"[PyPDF2] Error: {str(e)}")
            raise
        
        return text, metadata
    
    def extract_with_pdfplumber(self, path: Path) -> Tuple[str, Dict]:
        """Extract using pdfplumber (accurate, handles tables)"""
        logger.info(f"[pdfplumber] Extracting: {path.name}")
        text = ""
        metadata = {'tables': [], 'num_pages': 0}
        
        try:
            with pdfplumber.open(path) as pdf:
                metadata['num_pages'] = len(pdf.pages)
                
                for page_num, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
                    
                    # Extract tables
                    tables = page.extract_tables()
                    if tables:
                        for table in tables:
                            if table:
                                table_text = "\n".join([
                                    " | ".join(str(cell) if cell else "" for cell in row) 
                                    for row in table
                                ])
                                text += f"\n[TABLE]\n{table_text}\n[/TABLE]\n\n"
                                metadata['tables'].append({
                                    'page': page_num + 1,
                                    'rows': len(table)
                                })
                
                logger.info(f"[pdfplumber] Extracted {len(text)} chars, {len(metadata['tables'])} tables")
                            
        except Exception as e:
            logger.error(f"[pdfplumber] Error: {str(e)}")
            raise
        
        return text, metadata
    
    def extract_with_ocr(self, path: Path, language: str = 'fra+eng') -> Tuple[str, Dict]:
        """Extract using OCR (for scanned PDFs)"""
        logger.info(f"[OCR] Extracting: {path.name}")
        text = ""
        metadata = {'ocr_used': True, 'ocr_language': language}
        
        try:
            logger.info("[OCR] Converting PDF to images...")
            images = convert_from_path(path, dpi=300)
            metadata['num_pages'] = len(images)
            
            for i, image in enumerate(images):
                logger.info(f"[OCR] Processing page {i+1}/{len(images)}")
                page_text = pytesseract.image_to_string(image, lang=language)
                text += page_text + "\n\n"
            
            logger.info(f"[OCR] Extracted {len(text)} chars via OCR")
                
        except Exception as e:
            logger.error(f"[OCR] Error: {str(e)}")
            logger.error("Check if Tesseract is installed: apt-get install tesseract-ocr")
            raise
        
        return text, metadata
    
    def extract_hybrid(self, path: Path) -> Tuple[str, Dict]:
        """Hybrid method: try pdfplumber, fallback to OCR if needed"""
        logger.info(f"[HYBRID] Extracting: {path.name}")
        
        try:
            text, metadata = self.extract_with_pdfplumber(path)
            
            if len(text.strip()) < 100:
                logger.warning("[HYBRID] Low text extraction, switching to OCR")
                text, ocr_metadata = self.extract_with_ocr(path)
                metadata.update(ocr_metadata)
            else:
                logger.info("[HYBRID] Extraction successful with pdfplumber")
                
        except Exception as e:
            logger.warning(f"[HYBRID] pdfplumber failed: {str(e)}, trying OCR...")
            try:
                text, metadata = self.extract_with_ocr(path)
            except Exception as ocr_error:
                logger.error(f"[HYBRID] OCR also failed: {str(ocr_error)}")
                logger.warning("[HYBRID] Last attempt with PyPDF2...")
                text, metadata = self.extract_with_pypdf2(path)
        
        return text, metadata
    
    def extract(self, 
                path: Path, 
                method: PDFExtractionMethod = PDFExtractionMethod.HYBRID,
                clean_text: bool = True) -> Tuple[str, Dict]:
        """
        Main extraction entry point
        
        Args:
            path: Path to PDF file
            method: Extraction method to use
            clean_text: Apply automatic text cleaning
        
        Returns:
            Tuple of (extracted_text, metadata)
        """
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")
        
        if path.suffix.lower() != '.pdf':
            raise ValueError(f"File must be PDF: {path}")
        
        logger.info(f"=== Starting extraction: {path.name} ===")
        
        if method == PDFExtractionMethod.PYPDF2:
            text, metadata = self.extract_with_pypdf2(path)
        elif method == PDFExtractionMethod.PDFPLUMBER:
            text, metadata = self.extract_with_pdfplumber(path)
        elif method == PDFExtractionMethod.OCR:
            text, metadata = self.extract_with_ocr(path)
        else:  # HYBRID
            text, metadata = self.extract_hybrid(path)
        
        if clean_text and text:
            text = self.cleaner.clean(text)
        
        logger.info(f"=== Extraction complete: {len(text)} characters ===")
        return text, metadata


class Document:
    """Represents a loaded document with content and metadata"""
    
    def __init__(
        self,
        content: str,
        metadata: Optional[Dict] = None,
        doc_id: Optional[str] = None
    ):
        self.content = content
        self.metadata = metadata or {}
        self.doc_id = doc_id or self._generate_id()
    
    def _generate_id(self) -> str:
        """Generate a unique document ID"""
        return f"doc_{datetime.now().strftime('%Y%m%d_%H%M%S_%f')}"
    
    def __repr__(self) -> str:
        return f"Document(id={self.doc_id}, length={len(self.content)})"


class DocumentLoader:
    """Enhanced document loader with robust PDF processing"""
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.txt', '.docx'}
    
    def __init__(self, 
                 pdf_method: PDFExtractionMethod = PDFExtractionMethod.HYBRID,
                 aggressive_cleaning: bool = False):
        """
        Initialize document loader
        
        Args:
            pdf_method: PDF extraction method (pypdf2, pdfplumber, ocr, hybrid)
            aggressive_cleaning: Enable aggressive text cleaning
        """
        self.loaded_documents: Dict[str, Document] = {}
        self.pdf_method = pdf_method
        self.text_cleaner = TextCleaner(aggressive=aggressive_cleaning)
        self.pdf_extractor = PDFExtractor(cleaner=self.text_cleaner)
        
        logger.info(f"DocumentLoader initialized with method: {pdf_method.value}")
    
    def detect_language(self, text: str) -> Optional[str]:
        """Detect document language"""
        try:
            sample = text[:2000]
            lang = detect(sample)
            logger.info(f"Language detected: {lang}")
            return lang
        except LangDetectException:
            logger.warning("Could not detect language")
            return None
    
    def load(self, file_path: str) -> Document:
        """
        Load a document from file path with enhanced PDF support
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Document object with content and metadata
            
        Raises:
            ValueError: If file format is not supported
            FileNotFoundError: If file does not exist
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        extension = path.suffix.lower()
        
        if extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file format: {extension}. "
                f"Supported formats: {self.SUPPORTED_EXTENSIONS}"
            )
        
        logger.info(f"Loading document: {path.name}")
        
        # Extract content based on file type
        if extension == '.pdf':
            content, pdf_metadata = self._load_pdf_enhanced(path)
            extra_metadata = pdf_metadata
        elif extension == '.txt':
            content = self._load_txt(path)
            extra_metadata = {}
        elif extension == '.docx':
            content = self._load_docx(path)
            extra_metadata = {}
        else:
            raise ValueError(f"Unsupported extension: {extension}")
        
        # Detect language
        language = self.detect_language(content)
        
        # Create metadata
        metadata = {
            'filename': path.name,
            'filepath': str(path.absolute()),
            'extension': extension,
            'size_bytes': path.stat().st_size,
            'loaded_at': datetime.now().isoformat(),
            'language': language,
            'char_count': len(content),
            'word_count': len(content.split()),
            **extra_metadata
        }
        
        # Create document
        document = Document(content=content, metadata=metadata)
        self.loaded_documents[document.doc_id] = document
        
        logger.info(f"✅ Document loaded: {path.name} ({len(content)} chars)")
        return document
    
    def _load_pdf_enhanced(self, path: Path) -> Tuple[str, Dict]:
        """Load PDF with enhanced extraction"""
        content, metadata = self.pdf_extractor.extract(
            path, 
            method=self.pdf_method,
            clean_text=True
        )
        
        if not content or len(content.strip()) < 50:
            raise ValueError(f"Insufficient text extracted (<50 chars): {len(content)}")
        
        return content, metadata
    
    def _load_txt(self, path: Path) -> str:
        """Load content from TXT file"""
        with open(path, 'r', encoding='utf-8') as file:
            content = file.read()
        
        # Apply basic cleaning
        content = self.text_cleaner.clean(content)
        return content
    
    def _load_docx(self, path: Path) -> str:
        """Load content from DOCX file"""
        try:
            from docx import Document as DocxDocument
            
            doc = DocxDocument(path)
            content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text)
            
            text = '\n\n'.join(content)
            
            # Apply cleaning
            text = self.text_cleaner.clean(text)
            return text
            
        except ImportError:
            raise ImportError(
                "python-docx is required for DOCX support. "
                "Install it with: pip install python-docx"
            )
    
    def load_multiple(self, file_paths: List[str]) -> List[Document]:
        """
        Load multiple documents
        
        Args:
            file_paths: List of file paths
            
        Returns:
            List of Document objects
        """
        documents = []
        
        logger.info(f"Loading {len(file_paths)} documents...")
        
        for i, file_path in enumerate(file_paths, 1):
            try:
                logger.info(f"[{i}/{len(file_paths)}] {Path(file_path).name}")
                doc = self.load(file_path)
                documents.append(doc)
            except Exception as e:
                logger.error(f"❌ Error loading {file_path}: {e}")
        
        logger.info(f"✅ Loaded {len(documents)}/{len(file_paths)} documents")
        return documents
    
    def load_directory(self, directory_path: str, pattern: str = "*") -> List[Document]:
        """
        Load all supported documents from a directory
        
        Args:
            directory_path: Path to directory
            pattern: File pattern (e.g., "*.pdf")
        
        Returns:
            List of loaded documents
        """
        directory = Path(directory_path)
        
        if not directory.exists():
            raise FileNotFoundError(f"Directory not found: {directory_path}")
        
        # Find all supported files
        file_paths = []
        for ext in self.SUPPORTED_EXTENSIONS:
            file_paths.extend(directory.glob(f"{pattern}{ext}"))
        
        logger.info(f"Found {len(file_paths)} files in {directory_path}")
        
        return self.load_multiple([str(p) for p in file_paths])
    
    def get_document(self, doc_id: str) -> Optional[Document]:
        """Get a loaded document by ID"""
        return self.loaded_documents.get(doc_id)
    
    def list_documents(self) -> List[Dict]:
        """List all loaded documents with metadata"""
        return [
            {
                'doc_id': doc_id,
                'filename': doc.metadata.get('filename'),
                'extension': doc.metadata.get('extension'),
                'size': len(doc.content),
                'language': doc.metadata.get('language'),
                'loaded_at': doc.metadata.get('loaded_at'),
                'extraction_method': doc.metadata.get('extraction_method', 'standard')
            }
            for doc_id, doc in self.loaded_documents.items()
        ]
    
    def get_statistics(self) -> Dict:
        """Get loader statistics"""
        if not self.loaded_documents:
            return {'total_documents': 0}
        
        total_chars = sum(len(doc.content) for doc in self.loaded_documents.values())
        
        extensions = {}
        languages = {}
        
        for doc in self.loaded_documents.values():
            ext = doc.metadata.get('extension', 'unknown')
            lang = doc.metadata.get('language', 'unknown')
            
            extensions[ext] = extensions.get(ext, 0) + 1
            languages[lang] = languages.get(lang, 0) + 1
        
        return {
            'total_documents': len(self.loaded_documents),
            'total_characters': total_chars,
            'avg_chars_per_doc': total_chars // len(self.loaded_documents),
            'by_extension': extensions,
            'by_language': languages
        }
    
    def clear(self):
        """Clear all loaded documents"""
        self.loaded_documents.clear()
        logger.info("All documents cleared")


# ============================================
# USAGE EXAMPLES
# ============================================

if __name__ == "__main__":
    """Examples of using the enhanced document loader"""
    
    # Example 1: Load a single PDF with hybrid method
    print("\n=== Example 1: Load PDF with hybrid method ===")
    loader = DocumentLoader(
        pdf_method=PDFExtractionMethod.HYBRID,
        aggressive_cleaning=False
    )
    
    try:
        doc = loader.load("C:\\Users\\lenovo\\OneDrive\\Bureau\\RAG System\\data\\code de travail.pdf")
        print(f"✅ Loaded: {doc.metadata['filename']}")
        print(f"📊 Length: {len(doc.content)} chars")
        print(f"🌍 Language: {doc.metadata.get('language')}")
        print(f"📄 Pages: {doc.metadata.get('num_pages')}")
        print(f"\n📝 Preview:\n{doc.content[:200]}...")
    except FileNotFoundError:
        print("❌ example.pdf not found")
    
    # Example 4: Get statistics
    print("\n=== Example 4: Statistics ===")
    stats = loader.get_statistics()
    print(f"Total documents: {stats.get('total_documents', 0)}")
    print(f"By extension: {stats.get('by_extension', {})}")
    print(f"By language: {stats.get('by_language', {})}")
    
    # Example 5: List all documents
    print("\n=== Example 5: List documents ===")
    for doc_info in loader.list_documents():
        print(f"  - {doc_info['filename']} ({doc_info['extension']}) - {doc_info['language']}")